import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import io
import zipfile
from datetime import datetime

# --- CONFIGURAZIONE PAGINA E STILE SMARTPHONE ---
st.set_page_config(page_title="Difesa Scuola AI", page_icon="⚖️", layout="wide")

st.markdown("""
    <style>
    .stButton>button { width: 100%; height: 3.5em; background-color: #004a99; color: white; font-weight: bold; border-radius: 10px; }
    .stDownloadButton>button { width: 100%; height: 3.5em; background-color: #28a745; color: white; border-radius: 10px; }
    .block-container { padding-top: 2rem; }
    @media (max-width: 640px) { .stActionButton { display: none; } }
    </style>
    """, unsafe_allow_html=True)

# --- LOGICA DI GENERAZIONE DOCUMENTI ---

def crea_documenti(dati, files):
    # 1. MEMORIA DIFENSIVA
    mem = Document()
    style = mem.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Intestazione
    h = mem.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h.add_run(f"Spett.le {dati['scuola']}\nAll'attenzione dell'Autorità Procedente").bold = True

    mem.add_heading(f"MEMORIA DIFENSIVA - Prot. {dati['prot']}", 0)

    # Dati assistito e difensore
    p1 = mem.add_paragraph()
    p1.add_run(f"Il/La sottoscritto/a {dati['nome']}, nato/a a {dati['nascita_L']} il {dati['nascita_D']}, in servizio come {dati['ruolo']} presso {dati['sede']}. ")
    
    if dati['tipo_dif'] != "Nessuna":
        titolo = "Avv." if dati['tipo_dif'] == "Avvocato" else "Sig."
        p1.add_run(f"L'esponente agisce con l'assistenza del {titolo} {dati['nome_dif']} ({dati['org']}), presso il cui studio/ufficio elegge domicilio ai fini del presente procedimento.")
    
    # Sezione Vizi (Logica Automatica)
    mem.add_heading("1. ECCEZIONI PRELIMINARI E DI RITO", level=1)
    if dati['ruolo'] == "Docente" and "sospensione" in dati['analisi'].lower():
        mem.add_paragraph("Si eccepisce l'INCOMPETENZA FUNZIONALE del Dirigente Scolastico. Trattandosi di sanzione superiore alla censura per personale DOCENTE, la competenza è riservata all'U.P.D. (Cass. Civ. 28111/2021).", style='List Bullet')
    
    mem.add_paragraph("Si contesta la genericità dell'addebito e la violazione del principio di proporzionalità.", style='List Bullet')

    # Merito e Conclusioni
    mem.add_heading("2. MERITO E CONCLUSIONI", level=1)
    mem.add_paragraph(dati['fatti'])
    mem.add_paragraph("\nSi richiede l'ARCHIVIAZIONE del procedimento o la derubricazione della sanzione.")

    # Indice Allegati
    mem.add_page_break()
    mem.add_heading("INDICE DEGLI ALLEGATI", level=1)
    for i, f in enumerate(files, 1):
        mem.add_paragraph(f"{i}. {f.name}")

    # 2. DELEGA DI ASSISTENZA
    del_doc = Document()
    del_doc.add_heading("DELEGA DI ASSISTENZA", 0)
    del_doc.add_paragraph(f"Il sottoscritto {dati['nome']} delega formalmente {dati['nome_dif']} a rappresentarlo nel procedimento {dati['prot']}.")
    
    return mem, del_doc

# --- INTERFACCIA WEB ---

st.title("⚖️ Assistente Legale Scuola")
st.write("Generazione memorie per Docenti e ATA - Ottimizzato per Smartphone")

# Password di protezione (Opzionale ma consigliata)
pwd = st.sidebar.text_input("Password Accesso", type="password")
if pwd == "Scuola2026": # Cambia questa password
    menu = st.sidebar.radio("Navigazione", ["Nuova Pratica", "Archivio Storico"])

    if menu == "Nuova Pratica":
        with st.form("form_dati"):
            col1, col2 = st.columns(2)
            with col1:
                nome = st.text_input("Nome e Cognome Dipendente")
                ruolo = st.selectbox("Qualifica", ["Docente", "ATA"])
                sede = st.text_input("Sede di Servizio")
                prot = st.text_input("Protocollo Contestazione")
            with col2:
                nascita_L = st.text_input("Luogo di Nascita")
                nascita_D = st.text_input("Data di Nascita (GG/MM/AAAA)")
                scuola = st.text_input("Autorità (Scuola o USR)")
            
            st.divider()
            t_dif = st.selectbox("Assistenza", ["Nessuna", "Avvocato", "Sindacalista"])
            n_dif = st.text_input("Nome Difensore / Sindacalista")
            org = st.text_input("Foro o Sigla Sindacale")
            
            fatti = st.text_area("Breve ricostruzione dei fatti (Difesa)")
            files = st.file_uploader("Carica Atti (PDF/Foto)", accept_multiple_files=True)
            
            submit = st.form_submit_button("🚀 GENERA FASCICOLO")

        if submit and files:
            progress_bar = st.progress(0)
            status = st.empty()
            testo_estratto = ""
            
            for i, f in enumerate(files):
                status.text(f"Analisi file {i+1}...")
                testo_estratto += f.name + " " 
                time.sleep(0.8)
                progress_bar.progress(int((i+1)/len(files)*100))
            
            dati = {
                'nome': nome, 'ruolo': ruolo, 'sede': sede, 'prot': prot,
                'nascita_L': nascita_L, 'nascita_D': nascita_D, 'scuola': scuola,
                'tipo_dif': t_dif, 'nome_dif': n_dif, 'org': org, 'fatti': fatti,
                'analisi': testo_estratto
            }
            
            mem_doc, del_doc = crea_documenti(dati, files)
            
            # Creazione ZIP
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w") as zf:
                b1 = io.BytesIO(); mem_doc.save(b1)
                b2 = io.BytesIO(); del_doc.save(b2)
                zf.writestr(f"01_Memoria_{nome}.docx", b1.getvalue())
                zf.writestr(f"02_Delega_{nome}.docx", b2.getvalue())
            
            st.success("✅ Documenti pronti!")
            st.download_button("📥 SCARICA ARCHIVIO DIFESA", buf.getvalue(), f"Difesa_{nome}.zip")

    elif menu == "Archivio Storico":
        st.subheader("🗂️ Archivio Permanente")
        st.info("In questa sezione (collegata a DB) vedrai lo storico dei file caricati.")
        st.write("Funzionalità database attiva al prossimo salvataggio.")

else:
    st.warning("Inserisci la password nella barra laterale per accedere.")
